# import the necessary libraries
import argparse
import os
import sys
import whisperx
import torch
from docx import Document
from dotenv import load_dotenv
from huggingface_hub import whoami
from whisperx.utils import get_writer

# Load environment variables from a .env file (Hugging Face API tokens)
load_dotenv()

# Helper function to format timestamps into a human-readable format
def format_timestamp(seconds: float, always_include_hours: bool = False, decimal_marker: str = ".") -> str:
    assert seconds >= 0, "Non-negative timestamp expected"
    milliseconds = round(seconds * 1000.0)

    hours = milliseconds // 3_600_000
    milliseconds -= hours * 3_600_000

    minutes = milliseconds // 60_000
    milliseconds -= minutes * 60_000

    seconds = milliseconds // 1_000
    milliseconds -= seconds * 1_000

    hours_marker = f"{hours:02d}:" if always_include_hours or hours > 0 else ""
    return f"{hours_marker}{minutes:02d}:{seconds:02d}{decimal_marker}{milliseconds:03d}"

# Main function to process audio files for transcription
def main(input_dirs, language="en", max_speakers=5, device_id=0, device='unspecified'):
    # Authenticate with Hugging Face 
    info = whoami(os.getenv("HF_TOKEN"))
    
    # Automatically select the appropriate device (GPU if available, otherwise CPU)
    if device == 'unspecified':
        device = "cuda" if torch.cuda.is_available() else "cpu"
    elif device=="cuda" or device=="cpu":
        pass
    else:
        print("unsupported device")
        return
    print(f"Using device: {device}")
    
    # Set the compute precision type (float32 for compatibility across devices)
    compute_type = "float32"
    model_names = ["large-v3"]  # Specify the transcription model to use

    # Iterate through each specified language and input directory
    for audio_language in [language]:
        for input_dir in input_dirs:
            # Create an output directory to store transcription results
            output_dir = os.path.join(input_dir, "output")
            os.makedirs(output_dir, exist_ok=True)

            # Collect all audio files (supported formats: .wav, .mp3, .m4a)
            wav_files = [
                f for f in os.listdir(input_dir)
                if f.lower().endswith((".wav", ".mp3", ".m4a"))
            ]

            # Loop through each model (only large-v3 in this case)
            for model_name in model_names:
                # Load the WhisperX model for transcription
                model = whisperx.load_model(
                    model_name,
                    device,
                    device_index=device_id,
                    compute_type=compute_type,
                    language=audio_language,
                )

                # Process each audio file in the directory
                for file in wav_files:
                    # Define paths for output files (text and Word document)
                    output_txt = os.path.join(output_dir, f"whisperX-{file}-{model_name}.txt")
                    output_docx = os.path.join(output_dir, f"whisperX-{file}-{model_name}.docx")

                    # Skip files that have already been processed
                    if os.path.exists(output_txt):
                        print(f"Skipping {file} (already processed)")
                        continue

                    # Full path to the current audio file
                    full_path = os.path.join(input_dir, file)
                    print(f"Processing: {full_path}")

                    # Load the audio and prepare it for transcription
                    audio = whisperx.load_audio(full_path)
                    
                    # Transcribe the audio using the WhisperX model
                    result = model.transcribe(audio, batch_size=batch_size, language=audio_language)

                    # Align the transcription with the audio
                    align_model, metadata = whisperx.load_align_model(language_code=audio_language, device=device)
                    result = whisperx.align(result["segments"], align_model, metadata, audio, device)

                    # Perform speaker diarization to assign speakers to segments
                    diarize_model = whisperx.DiarizationPipeline(use_auth_token=os.getenv("HF_TOKEN"), device=device)
                    diarize_segments = diarize_model(audio, min_speakers=1, max_speakers=max_speakers)
                    result = whisperx.assign_word_speakers(diarize_segments, result)

                    # Write the transcription to a plain text file with timestamps and speaker labels
                    with open(output_txt, "w", encoding="utf-8") as txt_file:
                        current_speaker = None
                        for segment in result["segments"]:
                            start_time = format_timestamp(segment["start"])
                    
                            if "speaker" in segment and segment["speaker"] != current_speaker:
                                current_speaker = segment["speaker"]
                                # Write start time and speaker
                                txt_file.write(f"\n\n[{start_time}] <|{current_speaker}|>: ")
                            txt_file.write(segment["text"].strip() + " ")

                    # Write the transcription to a Word document with formatted paragraphs
                    document = Document()
                    current_speaker = None
                    paragraph = ""
                    for segment in result["segments"]:
                        start_time = format_timestamp(segment["start"])
                    
                        if "speaker" in segment and segment["speaker"] != current_speaker:
                            if paragraph:
                                document.add_paragraph(paragraph)
                            current_speaker = segment["speaker"]
                            # Add start time and speaker
                            paragraph = f"[{start_time}] <|{current_speaker}|>: "
                        paragraph += segment["text"].strip() + " "
                    if paragraph:
                        document.add_paragraph(paragraph)
                    document.save(output_docx)

                    # Add language information to the transcription result
                    result['language'] = language
                    
                    # Save transcription in additional formats (JSON, TXT, SRT)
                    whisperx_writer_dir = os.path.join(output_dir, model_name)
                    os.makedirs(whisperx_writer_dir, exist_ok=True)
                    writer_options = {
                        "max_line_width": None,
                        "max_line_count": None,
                        "highlight_words": False,
                    }
                    for fmt in ["json", "txt", "srt"]:
                        writer = get_writer(fmt, whisperx_writer_dir)
                        writer(result, file, writer_options)

# Entry point for the script, handling command-line arguments
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="WhisperX Audio Transcription Tool")
    parser.add_argument(
        "input_dirs", nargs="+",
        help="Directories containing audio files for transcription"
    )
    parser.add_argument(
        "-l", "--language", type=str, default="en",
        help="Language for transcription (default: 'en')"
    )
    parser.add_argument(
        "-ms", "--max_speakers", type=int, default=5,
        help="Maximum number of speakers to identify (default: 5)"
    )
    parser.add_argument(
        "-g", "--gpu", type=int, default=0,
        help="GPU ID to use (default: 0)"
    )
    parser.add_argument(
        "-d", "--device", type=str, default="unspecified",
        help="device to use, CUDA or cpu (default: unspecified)"
    )
    args = parser.parse_args()
    main(args.input_dirs, args.language, args.max_speakers, args.gpu)

